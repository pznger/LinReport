import docx
import sys
import os
from docx import Document
from docx.shared import Inches
from PyQt5.QtWidgets import QApplication, QWidget, QPushButton, QFileDialog, QVBoxLayout, QLabel, QLineEdit
from openai import OpenAI

def generate_report(java_project_path, output_doc_path, requirements_question, image_class_diagram_path, image_result_path):
    # 完善信息
    requirements_question = requirements_question + ",给出这个实验的实验目的，三条左右，不要带有任何格式，每条十几个字即可，不要加标题。"
    summary_question = requirements_question + "，给出这个实验的实验总结，300字左右，只给出一段话即可不带有任何格式，不要加标题，这个回答中不要带有实验目的了，只含有实验报告，并且只包含实验总结，不举具体的例子。"

    # 定义api_key
    api_key = "sk-JGR1yG7xV86kqxDcty03hwEAWhXhaVnwcxEKGU0Fp40rAbgs"
    base_url = "https://api.moonshot.cn/v1"

    # 创建Document对象
    doc = Document()

    # 设置字体加粗
    bold = True

    # 调用Kimi API接口获取实验内容和要求
    client = OpenAI(
        api_key=api_key,
        base_url=base_url
    )

    # 获取实验内容和要求
    requirements_completion = client.chat.completions.create(
        model="moonshot-v1-8k",
        messages=[
            {
                "role": "system",
                "content": "你是 Kimi，由 Moonshot AI 提供的人工智能助手..."
            },
            {
                "role": "user",
                "content": requirements_question
            },
        ],
        temperature=0.3,
    )

    requirements_answer = requirements_completion.choices[0].message.content

    # 获取实验总结与感想
    summary_completion = client.chat.completions.create(
        model="moonshot-v1-8k",
        messages=[
            {
                "role": "system",
                "content": "你是 Kimi，由 Moonshot AI 提供的人工智能助手..."
            },
            {
                "role": "user",
                "content": summary_question
            },
        ],
        temperature=0.3,
    )

    summary_answer = summary_completion.choices[0].message.content

    # 添加实验内容和要求
    p = doc.add_paragraph('')
    p.add_run('一、实验内容和要求').bold = bold
    doc.add_paragraph(requirements_answer)

    # 添加实验内容和结果
    p = doc.add_paragraph('')
    p.add_run('二、实验内容和结果').bold = bold
    # 添加类图
    doc.add_paragraph('（1）类图')
    doc.add_picture(image_class_diagram_path, width=docx.shared.Inches(6))

    # 整合Java项目代码
    java_files_content = ""
    for root, dirs, files in os.walk(java_project_path):
        for file in files:
            if file.endswith('.java'):
                java_path = os.path.join(root, file)
                with open(java_path, 'r', encoding='utf-8') as f:
                    java_files_content += f'--- {file} from {root} ---\n'
                    java_files_content += f.read() + '\n\n'

    code_paragraph = doc.add_paragraph('')
    code_paragraph.add_run('（2）代码').bold = bold
    code_paragraph.add_run(java_files_content)

    # 添加运行结果
    code_paragraph = doc.add_paragraph('')
    code_paragraph.add_run('（3）运行结果').bold = bold
    doc.add_picture(image_result_path, width=docx.shared.Inches(6))

    # 添加实验总结与感想
    code_paragraph = doc.add_paragraph('')
    code_paragraph.add_run('三、实验总结及感想').bold = bold
    doc.add_paragraph(summary_answer)

    # 保存文档
    doc.save(output_doc_path)


def replace_single_slash(path):
    # 替换路径中的斜杠为双斜杠
    return path.replace('\\', '\\\\')


class MyApp(QWidget):
    def __init__(self):
        super().__init__()
        self.image_result_path_button = None
        self.label5 = None
        self.image_result_path_edit = None
        self.image_class_diagram_path_button = None
        self.image_class_diagram_path_edit = None
        self.label4 = None
        self.output_doc_path_button = None
        self.output_doc_path_edit = None
        self.label3 = None
        self.java_project_path_button = None
        self.java_project_path_edit = None
        self.label2 = None
        self.requirements_question_edit = None
        self.label1 = None
        self.run_button = None
        self.initUI()

    def initUI(self):
        self.setWindowTitle('文档和图片路径选择器')
        self.setGeometry(100, 100, 600, 400)

        layout = QVBoxLayout()

        self.label1 = QLabel('实验报告题目:', self)
        layout.addWidget(self.label1)
        self.requirements_question_edit = QLineEdit(self)
        layout.addWidget(self.requirements_question_edit)

        self.label2 = QLabel('Java项目路径:', self)
        layout.addWidget(self.label2)
        self.java_project_path_edit = QLineEdit(self)
        layout.addWidget(self.java_project_path_edit)
        self.java_project_path_button = QPushButton('选择路径', self)
        self.java_project_path_button.clicked.connect(self.get_java_project_path)
        layout.addWidget(self.java_project_path_button)

        self.label3 = QLabel('文档输出路径:', self)
        layout.addWidget(self.label3)
        self.output_doc_path_edit = QLineEdit(self)
        layout.addWidget(self.output_doc_path_edit)
        self.output_doc_path_button = QPushButton('选择路径', self)
        self.output_doc_path_button.clicked.connect(self.get_output_doc_path)
        layout.addWidget(self.output_doc_path_button)

        self.label4 = QLabel('类图路径:', self)
        layout.addWidget(self.label4)
        self.image_class_diagram_path_edit = QLineEdit(self)
        layout.addWidget(self.image_class_diagram_path_edit)
        self.image_class_diagram_path_button = QPushButton('选择路径', self)
        self.image_class_diagram_path_button.clicked.connect(self.get_image_class_diagram_path)
        layout.addWidget(self.image_class_diagram_path_button)

        self.label5 = QLabel('运行结果图片路径:', self)
        layout.addWidget(self.label5)
        self.image_result_path_edit = QLineEdit(self)
        layout.addWidget(self.image_result_path_edit)
        self.image_result_path_button = QPushButton('选择路径', self)
        self.image_result_path_button.clicked.connect(self.get_image_result_path)
        layout.addWidget(self.image_result_path_button)

        self.run_button = QPushButton('生成实验报告', self)
        self.run_button.clicked.connect(self.run_script)
        layout.addWidget(self.run_button)

        self.setLayout(layout)

    def get_java_project_path(self):
        path = QFileDialog.getExistingDirectory(self, '选择Java项目路径')
        self.java_project_path_edit.setText(path.replace('/', '\\'))

    def get_output_doc_path(self):
        path, _ = QFileDialog.getSaveFileName(self, '选择文档输出路径', '', '*.docx')
        self.output_doc_path_edit.setText(path.replace('/', '\\'))

    def get_image_class_diagram_path(self):
        path, _ = QFileDialog.getOpenFileName(self, '选择类图路径', '', '*.png')
        self.image_class_diagram_path_edit.setText(path.replace('/', '\\'))

    def get_image_result_path(self):
        path, _ = QFileDialog.getOpenFileName(self, '选择运行结果图片路径', '', '*.png')
        self.image_result_path_edit.setText(path.replace('/', '\\'))

    def run_script(self):
        # 获取用户输入的路径和题目
        requirements_question = self.requirements_question_edit.text()
        java_project_path = self.java_project_path_edit.text()
        output_doc_path = self.output_doc_path_edit.text()
        image_class_diagram_path = self.image_class_diagram_path_edit.text()
        image_result_path = self.image_result_path_edit.text()

        # 替换路径中的斜杠为双斜杠
        java_project_path = replace_single_slash(java_project_path)
        output_doc_path = replace_single_slash(output_doc_path)
        image_class_diagram_path = replace_single_slash(image_class_diagram_path)
        image_result_path = replace_single_slash(image_result_path)

        # 这里调用你的脚本逻辑，使用上面获取的路径和题目
        print(f"实验报告题目: {requirements_question}")
        print(f"Java项目路径: {java_project_path}")
        print(f"文档输出路径: {output_doc_path}")
        print(f"类图路径: {image_class_diagram_path}")
        print(f"运行结果图片路径: {image_result_path}")

        self.statusBar().showMessage(f"报告已生成并保存在：{output_doc_path}")
        generate_report(java_project_path,output_doc_path,requirements_question,image_class_diagram_path,image_result_path)

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = MyApp()
    ex.show()
    sys.exit(app.exec_())