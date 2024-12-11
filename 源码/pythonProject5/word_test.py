from docx import Document

# 创建Document对象，这将创建一个新的Word文档
doc = Document()

# 向文档添加一个段落
doc.add_paragraph('Hello, this is a paragraph.')

# 添加另一个段落，并设置字体加粗
p = doc.add_paragraph('This is another paragraph.')
run = p.add_run(' And this is a run with bold text.')
run.bold = True

# 保存文档
doc.save('example.docx')