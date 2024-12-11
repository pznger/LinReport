import docx
from docx import Document
from openai import OpenAI
import os

# 用户输入的信息
java_project_path = "G:\\huidu\\软件设计实验报告\\模板模式\\barista"    #Java项目路径
output_doc_path = "G:\\huidu\\软件设计实验报告\\test.docx"   #文档输出路径
requirements_question = "使用模版方法模式对咖啡因饮料加工模拟案例进行建模和实现"    #实验标题
image_class_diagram_path = "G:\\huidu\\软件设计实验报告\\模板模式\\2.png"    #类图路径
image_result_path = "G:\\huidu\\软件设计实验报告\\模板模式\\2.png"  #运行结果图片路径

# 完善信息
requirements_question = requirements_question + ",给出这个实验的实验目的，三条左右，不要带有任何格式，每条十几个字即可，不要加标题。"
summary_question = requirements_question + "，给出这个实验的实验总结，300字左右，只给出一段话即可不带有任何格式，不要加标题，这个回答中不要带有实验目的了，只含有实验报告，并且只包含实验总结，不举具体的例子。"

# 定义api_key
api_key = "sk-JGR1yG7xV86kqxDcty03hwEAWhXhaVnwcxEKGU0Fp40rAbgs"
base_url = "https://api.moonshot.cn/v1"

# 创建Document对象
doc = Document()

# 设置字体加粗
bold = True

# 调用Kimi API接口获取实验内容和要求
client = OpenAI(
    api_key=api_key,
    base_url=base_url
)

# 获取实验内容和要求
requirements_completion = client.chat.completions.create(
    model="moonshot-v1-8k",
    messages=[
        {
            "role": "system",
            "content": "你是 Kimi，由 Moonshot AI 提供的人工智能助手..."
        },
        {
            "role": "user",
            "content": requirements_question
        },
    ],
    temperature=0.3,
)

requirements_answer = requirements_completion.choices[0].message.content

# 获取实验总结与感想
summary_completion = client.chat.completions.create(
    model="moonshot-v1-8k",
    messages=[
        {
            "role": "system",
            "content": "你是 Kimi，由 Moonshot AI 提供的人工智能助手..."
        },
        {
            "role": "user",
            "content": summary_question
        },
    ],
    temperature=0.3,
)

summary_answer = summary_completion.choices[0].message.content

# 添加实验内容和要求
p = doc.add_paragraph('')
p.add_run('一、实验内容和要求').bold = bold
doc.add_paragraph(requirements_answer)

# 添加实验内容和结果
p = doc.add_paragraph('')
p.add_run('二、实验内容和结果').bold = bold
# 添加类图
doc.add_paragraph('（1）类图')
doc.add_picture(image_class_diagram_path, width=docx.shared.Inches(6))

# 整合Java项目代码
java_files_content = ""
for root, dirs, files in os.walk(java_project_path):
    for file in files:
        if file.endswith('.java'):
            java_path = os.path.join(root, file)
            with open(java_path, 'r', encoding='utf-8') as f:
                java_files_content += f'--- {file} from {root} ---\n'
                java_files_content += f.read() + '\n\n'

code_paragraph = doc.add_paragraph('')
code_paragraph.add_run('（2）代码').bold = bold
code_paragraph.add_run(java_files_content)

# 添加运行结果
code_paragraph = doc.add_paragraph('')
code_paragraph.add_run('（3）运行结果').bold = bold
doc.add_picture(image_result_path, width=docx.shared.Inches(6))

# 添加实验总结与感想
code_paragraph = doc.add_paragraph('')
code_paragraph.add_run('三、实验总结及感想').bold = bold
doc.add_paragraph(summary_answer)

# 保存文档
doc.save(output_doc_path)